from pathlib import Path
from typing import Tuple


def is_list_line(line: str) -> bool:
    stripped = line.strip()
    if stripped.startswith("• ") or stripped.startswith("- "):
        return True
    if len(stripped) >= 3 and stripped[0].isdigit() and ") " in stripped[:4]:
        return True
    if len(stripped) >= 3 and stripped[0].isdigit() and ". " in stripped[:4]:
        return True
    if stripped.startswith("（") and "）" in stripped[:4]:
        return True
    return False


def split_line_style(line: str) -> Tuple[str, str]:
    stripped = line.rstrip()
    if not stripped.strip():
        return "blank", ""
    if is_list_line(stripped):
        return "list", stripped
    return "paragraph", stripped


def export_to_word(text: str, output_path: str) -> Path:
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except ImportError as error:
        raise RuntimeError("缺少 python-docx，请先安装：pip install python-docx") from error

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)
    try:
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    except Exception:
        pass

    for raw in text.split("\n"):
        line_type, line_value = split_line_style(raw)
        paragraph = document.add_paragraph()

        if line_type == "blank":
            paragraph.add_run("")
            continue

        if line_type == "list":
            paragraph.paragraph_format.left_indent = Pt(18)
            paragraph.paragraph_format.first_line_indent = Pt(-18)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.add_run(line_value)
            continue

        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.add_run(line_value)

    output = Path(output_path)
    document.save(output)
    return output


def export_to_pdf(text: str, output_path: str) -> Path:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as error:
        raise RuntimeError("缺少 reportlab，请先安装：pip install reportlab") from error

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    styles = getSampleStyleSheet()
    normal_style = ParagraphStyle(
        "NormalCN",
        parent=styles["Normal"],
        fontName="STSong-Light",
        fontSize=11,
        leading=18,
        spaceAfter=6,
    )
    list_style = ParagraphStyle(
        "ListCN",
        parent=normal_style,
        leftIndent=18,
        firstLineIndent=-12,
        spaceAfter=4,
    )

    story = []
    for raw in text.split("\n"):
        line_type, line_value = split_line_style(raw)
        if line_type == "blank":
            story.append(Spacer(1, 8))
            continue

        escaped = (
            line_value.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
        style = list_style if line_type == "list" else normal_style
        story.append(Paragraph(escaped, style))

    output = Path(output_path)
    document = SimpleDocTemplate(
        str(output),
        pagesize=A4,
        leftMargin=42,
        rightMargin=42,
        topMargin=42,
        bottomMargin=42,
    )
    document.build(story)
    return output
